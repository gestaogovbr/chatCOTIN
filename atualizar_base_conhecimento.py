#!/usr/bin/env python3
"""
Script para Atualização da Base de Conhecimento do ChatCOTIN
============================================================

Este script atualiza automaticamente a base de conhecimento do ChatCOTIN
quando novos documentos são adicionados à pasta 'Docs/'.

Formatos suportados:
- .docx (Microsoft Word)
- .md (Markdown)  
- .txt (Texto simples)

Uso:
    python atualizar_base_conhecimento.py

Autor: Sistema ChatCOTIN
Data: Janeiro 2025
"""

import os
import sys
import django
import shutil
import time
from docx import Document

def setup_django():
    """Configura o ambiente Django"""
    os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'core.settings')
    django.setup()

def extract_text_from_docx(file_path):
    """
    Extrai texto de arquivo Microsoft Word (.docx)
    
    Args:
        file_path (str): Caminho para o arquivo .docx
        
    Returns:
        LangchainDocument: Documento processado ou None se erro
    """
    try:
        print(f"  📄 Processando Word: {os.path.basename(file_path)}")
        doc = Document(file_path)
        
        text_content = []
        
        # Extrair parágrafos
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extrair tabelas
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        full_text = "\n".join(text_content)
        print(f"    ✅ Extraído: {len(full_text):,} caracteres")
        
        from langchain.schema import Document as LangchainDocument
        return LangchainDocument(
            page_content=full_text,
            metadata={
                "source": file_path,
                "filename": os.path.basename(file_path),
                "type": "word_document",
                "extension": ".docx"
            }
        )
        
    except Exception as e:
        print(f"    ❌ Erro ao processar {os.path.basename(file_path)}: {e}")
        return None

def extract_text_from_markdown(file_path):
    """
    Extrai texto de arquivo Markdown (.md)
    
    Args:
        file_path (str): Caminho para o arquivo .md
        
    Returns:
        LangchainDocument: Documento processado ou None se erro
    """
    try:
        print(f"  📝 Processando Markdown: {os.path.basename(file_path)}")
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        print(f"    ✅ Extraído: {len(content):,} caracteres")
        
        from langchain.schema import Document as LangchainDocument
        return LangchainDocument(
            page_content=content,
            metadata={
                "source": file_path,
                "filename": os.path.basename(file_path),
                "type": "markdown",
                "extension": ".md"
            }
        )
        
    except Exception as e:
        print(f"    ❌ Erro ao processar {os.path.basename(file_path)}: {e}")
        return None

def extract_text_from_txt(file_path):
    """
    Extrai texto de arquivo de texto simples (.txt)
    
    Args:
        file_path (str): Caminho para o arquivo .txt
        
    Returns:
        LangchainDocument: Documento processado ou None se erro
    """
    try:
        print(f"  📋 Processando Texto: {os.path.basename(file_path)}")
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        print(f"    ✅ Extraído: {len(content):,} caracteres")
        
        from langchain.schema import Document as LangchainDocument
        return LangchainDocument(
            page_content=content,
            metadata={
                "source": file_path,
                "filename": os.path.basename(file_path),
                "type": "plain_text",
                "extension": ".txt"
            }
        )
        
    except Exception as e:
        print(f"    ❌ Erro ao processar {os.path.basename(file_path)}: {e}")
        return None

def scan_documents_folder(docs_path):
    """
    Escaneia a pasta de documentos e identifica arquivos suportados
    
    Args:
        docs_path (str): Caminho para a pasta de documentos
        
    Returns:
        tuple: (docx_files, md_files, txt_files, total_files)
    """
    if not os.path.exists(docs_path):
        print(f"❌ Pasta {docs_path} não encontrada!")
        return [], [], [], 0
    
    all_files = os.listdir(docs_path)
    
    docx_files = [f for f in all_files if f.lower().endswith('.docx')]
    md_files = [f for f in all_files if f.lower().endswith('.md')]
    txt_files = [f for f in all_files if f.lower().endswith('.txt')]
    
    total_files = len(docx_files) + len(md_files) + len(txt_files)
    
    return docx_files, md_files, txt_files, total_files

def load_all_documents(docs_path):
    """
    Carrega todos os documentos suportados da pasta
    
    Args:
        docs_path (str): Caminho para a pasta de documentos
        
    Returns:
        list: Lista de documentos LangChain processados
    """
    docx_files, md_files, txt_files, total_files = scan_documents_folder(docs_path)
    
    print(f"📂 Documentos encontrados:")
    print(f"  📄 {len(docx_files)} arquivos Word (.docx)")
    print(f"  📝 {len(md_files)} arquivos Markdown (.md)")
    print(f"  📋 {len(txt_files)} arquivos Texto (.txt)")
    print(f"  📊 Total: {total_files} documentos")
    
    if total_files == 0:
        print("⚠️  Nenhum documento suportado encontrado!")
        return []
    
    print(f"\n🔄 Processando documentos...")
    all_docs = []
    
    # Processar arquivos Word
    for filename in docx_files:
        file_path = os.path.join(docs_path, filename)
        doc = extract_text_from_docx(file_path)
        if doc:
            all_docs.append(doc)
    
    # Processar arquivos Markdown
    for filename in md_files:
        file_path = os.path.join(docs_path, filename)
        doc = extract_text_from_markdown(file_path)
        if doc:
            all_docs.append(doc)
    
    # Processar arquivos Texto
    for filename in txt_files:
        file_path = os.path.join(docs_path, filename)
        doc = extract_text_from_txt(file_path)
        if doc:
            all_docs.append(doc)
    
    return all_docs

def create_knowledge_base(documents, collection_name, chroma_path):
    """
    Cria a base de conhecimento vetorial
    
    Args:
        documents (list): Lista de documentos processados
        collection_name (str): Nome da coleção ChromaDB
        chroma_path (str): Caminho para o banco ChromaDB
        
    Returns:
        Chroma: Instância do banco vetorial ou None se erro
    """
    if not documents:
        print("❌ Nenhum documento para processar!")
        return None
    
    try:
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        from langchain_community.embeddings import FastEmbedEmbeddings
        from langchain_community.vectorstores import Chroma
        
        # Estatísticas dos documentos
        total_chars = sum(len(doc.page_content) for doc in documents)
        print(f"\n📊 Estatísticas:")
        print(f"  📄 Documentos processados: {len(documents)}")
        print(f"  📝 Total de caracteres: {total_chars:,}")
        
        # Chunking otimizado
        print(f"\n✂️  Dividindo em chunks para processamento...")
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,          # Tamanho otimizado para português
            chunk_overlap=200,       # Overlap generoso para contexto
            separators=["\n\n", "\n", ". ", "! ", "? ", ", ", " "]
        )
        
        chunks = splitter.split_documents(documents)
        print(f"    ✅ Criados {len(chunks)} chunks otimizados")
        
        # Embeddings multilíngues otimizados para português
        print(f"\n🧠 Criando embeddings multilíngues...")
        embeddings = FastEmbedEmbeddings(
            model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
        )
        
        # Criar base vetorial
        print(f"💾 Salvando no ChromaDB...")
        vectorstore = Chroma.from_documents(
            chunks,
            embedding=embeddings,
            collection_name=collection_name,
            persist_directory=chroma_path
        )
        
        print(f"✅ Base de conhecimento criada com sucesso!")
        return vectorstore
        
    except Exception as e:
        print(f"❌ Erro ao criar base de conhecimento: {e}")
        import traceback
        traceback.print_exc()
        return None

def test_knowledge_base(vectorstore):
    """
    Testa a base de conhecimento com consultas específicas
    
    Args:
        vectorstore (Chroma): Instância do banco vetorial
    """
    print(f"\n🔍 Testando base de conhecimento...")
    
    test_queries = [
        "módulos da API de dados abertos",
        "transparência pública",
        "Lei de Acesso à Informação",
        "compras governamentais",
        "portal transparência",
        "dados de licitações"
    ]
    
    for query in test_queries:
        try:
            results = vectorstore.similarity_search(query, k=2)
            if results:
                print(f"  ✅ '{query}': {len(results)} resultado(s)")
                if results[0].metadata:
                    filename = results[0].metadata.get('filename', 'fonte desconhecida')
                    print(f"     📄 Fonte: {filename}")
            else:
                print(f"  ❌ '{query}': Nenhum resultado")
        except Exception as e:
            print(f"  ⚠️  '{query}': Erro no teste - {e}")

def main():
    """Função principal do script"""
    print("🚀 CHATCOTIN - Atualização da Base de Conhecimento")
    print("=" * 60)
    print("📋 Formatos suportados: .docx, .md, .txt")
    print("=" * 60)
    
    # Configurações
    docs_path = 'Docs'
    collection_name = f"chatcotin_knowledge_{int(time.time())}"
    chroma_path = "chroma_db"
    
    try:
        # 1. Configurar Django
        setup_django()
        
        # 2. Backup do banco antigo (se existir)
        if os.path.exists(chroma_path):
            try:
                backup_path = f"{chroma_path}_backup_{int(time.time())}"
                os.rename(chroma_path, backup_path)
                print(f"🗄️  Banco antigo movido para: {backup_path}")
            except Exception:
                print(f"⚠️  Banco antigo existe, usando nova coleção: {collection_name}")
        
        # 3. Carregar documentos
        documents = load_all_documents(docs_path)
        
        if not documents:
            print("\n❌ Nenhum documento foi carregado. Verifique a pasta 'Docs/'.")
            return False
        
        # 4. Criar base de conhecimento
        vectorstore = create_knowledge_base(documents, collection_name, chroma_path)
        
        if not vectorstore:
            print("\n❌ Falha ao criar base de conhecimento.")
            return False
        
        # 5. Testar base de conhecimento
        test_knowledge_base(vectorstore)
        
        # 6. Sucesso
        print(f"\n🎉 ATUALIZAÇÃO CONCLUÍDA COM SUCESSO!")
        print(f"   📄 {len(documents)} documentos indexados")
        print(f"   🔍 Base pronta para consultas do ChatCOTIN")
        print(f"   💾 Dados salvos em: {chroma_path}")
        
        return True
        
    except Exception as e:
        print(f"\n❌ Erro durante execução: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print(__doc__)
    success = main()
    sys.exit(0 if success else 1) 